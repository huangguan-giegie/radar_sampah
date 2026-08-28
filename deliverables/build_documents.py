from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_DIR = ROOT / "documents"
VERSION = "VERSION 2.0 | 15 AUGUST 2026"
BRAND = "TEAM 04 | DIVESAFE MY"
FOOTER_LABEL = f"{VERSION} | SYNTHETIC/PUBLIC DATA | NON-ENFORCEMENT DEMO"

TEAL = "0F6F73"
NAVY = "183B56"
INK_HEX = "243746"
MUTED = "61737C"
PALE = "EEF6F5"
PALE_BLUE = "EFF3F7"
GRID = "C9D6DB"
WHITE = "FFFFFF"
LINK = "0563C1"

INK = RGBColor.from_string(INK_HEX)


def set_run_font(run, name="Calibri", size=11, color=INK_HEX, bold=None, italic=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name, size, color, bold=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def set_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color, size=14, space=8):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def add_hyperlink(paragraph, label, url, size=10.5):
    rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    rpr.extend([fonts, color, underline, size_el])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([rpr, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|https?://\S+)"
)


def add_inline(paragraph, text, size=11, color=INK_HEX):
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(9, size - 1), color=NAVY)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            add_hyperlink(paragraph, link_match.group(1), link_match.group(2), size=size)
        else:
            url = token.rstrip(".,;")
            add_hyperlink(paragraph, url, url, size=size)
            if len(url) != len(token):
                run = paragraph.add_run(token[len(url):])
                set_run_font(run, size=size, color=color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def shade_cell(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, header=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    add_inline(paragraph, text, size=9.3, color=WHITE if header else INK_HEX)
    if header:
        for run in paragraph.runs:
            run.bold = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def table_widths(headers):
    lowered = [header.lower() for header in headers]
    count = len(headers)
    if count == 2:
        if lowered[0] == "field":
            return [2400, 6960]
        if "endpoint" in lowered[0]:
            return [2600, 6760]
        return [2700, 6660]
    if count == 3:
        if lowered[0].startswith("member"):
            return [2200, 2200, 4960]
        if lowered[0] == "value":
            return [1700, 2700, 4960]
        if lowered[0] == "stage" and lowered[1] == "action":
            return [900, 6500, 1960]
        if lowered[0] == "cadence":
            return [1800, 4260, 3300]
        if lowered[0] == "check":
            return [2250, 4010, 3100]
        return [1800, 4200, 3360]
    if count == 4:
        if lowered[1] == "date" and lowered[3] == "date":
            return [3100, 900, 4460, 900]
        if lowered[0] == "stage":
            return [1100, 3300, 1800, 3160]
        return [1800, 2400, 2400, 2760]
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def apply_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tblpr.find(qn(tag))
        if old is not None:
            tblpr.remove(old)
    tblw = OxmlElement("w:tblW")
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    tblind = OxmlElement("w:tblInd")
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.extend([tblw, tblind, layout])

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
        for index, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[index]))
            tcw.set(qn("w:type"), "dxa")


def add_table(doc, rows):
    widths = table_widths(rows[0])
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[index], value, header=True)
        shade_cell(table.rows[0].cells[index], NAVY)
    header_props = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_props.append(repeat)

    for row_index, row in enumerate(rows[1:], start=1):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if row_index % 2 == 1:
                shade_cell(cells[index], PALE_BLUE)
    apply_table_geometry(table, widths)
    table.rows[-1].cells[-1].paragraphs[-1].paragraph_format.space_after = Pt(0)


def create_abstract_numbering(doc, number_format, marker, marker_font):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), number_format)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), marker)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    ppr.extend([tabs, indent])
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), marker_font)
    fonts.set(qn("w:hAnsi"), marker_font)
    rpr.append(fonts)
    level.extend([start, fmt, text, justification, suffix, ppr, rpr])
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def create_num_instance(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, number])


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    apply_numbering(paragraph, num_id)
    add_inline(paragraph, text)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    add_inline(paragraph, text)


def add_metadata(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=9.5, color=NAVY, bold=True)
    add_inline(paragraph, value, size=9.5, color=MUTED)


def add_callout(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.10
    paragraph.paragraph_format.keep_together = True
    set_paragraph_shading(paragraph, PALE)
    set_paragraph_left_border(paragraph, TEAL)
    add_inline(paragraph, text, size=10.5)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=8, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, INK_HEX)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    set_style_font(title, "Calibri", 26, NAVY, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True

    for name, size, color, before, after in (
        ("Heading 1", 16, TEAL, 16, 8),
        ("Heading 2", 13, NAVY, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run(BRAND)
    set_run_font(header_run, size=8, color=TEAL, bold=True)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    footer_run = footer.add_run(FOOTER_LABEL)
    set_run_font(footer_run, size=8, color=MUTED)
    footer.add_run("\tPAGE ")
    for run in footer.runs[1:]:
        set_run_font(run, size=8, color=MUTED)
    add_page_field(footer)


def is_table_separator(line):
    if not line.startswith("|"):
        return False
    cells = [part.strip() for part in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        if not is_table_separator(line):
            rows.append([part.strip() for part in line.strip("|").split("|")])
        index += 1
    return rows, index


def parse_and_build(source):
    doc = Document()
    configure_document(doc)
    bullet_abstract = create_abstract_numbering(doc, "bullet", "•", "Symbol")
    checkbox_abstract = create_abstract_numbering(doc, "bullet", "☐", "Segoe UI Symbol")
    decimal_abstract = create_abstract_numbering(doc, "decimal", "%1.", "Calibri")
    bullet_num = create_num_instance(doc, bullet_abstract)
    checkbox_num = create_num_instance(doc, checkbox_abstract)
    decimal_num = create_num_instance(doc, decimal_abstract)

    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    first_title = True
    previous_kind = None
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            previous_kind = None
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1].strip()):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            previous_kind = "table"
            continue
        if line.startswith("# "):
            if first_title:
                kicker = doc.add_paragraph()
                kicker.paragraph_format.space_before = Pt(0)
                kicker.paragraph_format.space_after = Pt(3)
                kicker.paragraph_format.keep_with_next = True
                kicker_run = kicker.add_run(BRAND)
                set_run_font(kicker_run, size=9, color=TEAL, bold=True)
                paragraph = doc.add_paragraph(style="Title")
                title_text = line[2:].strip()
                paragraph.add_run(title_text)
                doc.core_properties.title = title_text
                doc.core_properties.subject = "Team 04 Radar Sampah MVP"
                doc.core_properties.author = "Team 04"
                doc.core_properties.keywords = "Radar Sampah, synthetic data, public sources, non-enforcement"
                first_title = False
            else:
                doc.add_heading(line[2:].strip(), level=1)
            previous_kind = "heading"
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            previous_kind = "heading"
            index += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            previous_kind = "heading"
            index += 1
            continue
        metadata = re.match(r"^\*\*(Document control|Audience|Format):\*\*\s*(.*)$", line)
        if metadata:
            add_metadata(doc, metadata.group(1), metadata.group(2))
            previous_kind = "metadata"
            index += 1
            continue
        if line.startswith("> "):
            add_callout(doc, line[2:].strip())
            previous_kind = "callout"
            index += 1
            continue
        checklist = re.match(r"^- \[[ xX]\]\s+(.*)$", line)
        if checklist:
            add_list_item(doc, checklist.group(1), checkbox_num)
            previous_kind = "checkbox"
            index += 1
            continue
        if line.startswith("- "):
            add_list_item(doc, line[2:].strip(), bullet_num)
            previous_kind = "bullet"
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if numbered:
            if numbered.group(1) == "1" and previous_kind != "decimal":
                decimal_num = create_num_instance(doc, decimal_abstract)
            add_list_item(doc, numbered.group(2), decimal_num)
            previous_kind = "decimal"
            index += 1
            continue
        if line.startswith("`") and line.endswith("`"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(line.strip("`").strip())
            set_run_font(run, name="Consolas", size=9.5, color=NAVY)
            set_paragraph_shading(paragraph, PALE_BLUE)
            previous_kind = "code"
            index += 1
            continue
        add_body(doc, line)
        previous_kind = "body"
        index += 1

    output = ROOT / f"{source.stem}.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    for markdown in sorted(SOURCE_DIR.glob("*.md")):
        print(parse_and_build(markdown))
